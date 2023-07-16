from docx import Document

def write_string_to_word(text, filename):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(filename)

def read_docx(file):
    doc = Document(file)
    text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
    return text